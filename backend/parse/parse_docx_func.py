import docx
import re

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fullText.append(para.text)
          
    return '\n'.join(fullText)

def parse_docx(filename):
    pattern1 = 'а-яА-ЯғқөңіәүұһҚҰҮӨҒӘІҢҒ.'
    pattern2 = 'А-ЯҚҰҮӨҒӘІҢҒ'

    doc = getText(f'./data/docx/{filename}')

    # with open('text.txt', 'w') as f:
    #     f.write(doc)

    x = re.findall(r"[{}][{} ]+[\n] [0-9 x\n]+[0-9][\n]".format(pattern2, pattern1), doc)

    result = {}

    for i in x:
        s = re.findall(r'[{}]+'.format(pattern1), i)
        n = re.findall(r'[0-9][0-9 ]+', i)
        s = " ".join(s)
        result[s] = n[0].replace(' ', '')

    return result


# print(x)